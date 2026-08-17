"""
Multi-model OCR service.

Strategy per file type:
  PDF (text-based)  → PyMuPDF direct text extraction (fast, accurate)
  PDF (scanned)     → PyMuPDF page render → Tesseract OCR per page
  Image             → Preprocessing pipeline → Tesseract
  Mixed/fallback    → Try text first, if <50 chars/page fall back to OCR

Image preprocessing chain (improves accuracy significantly):
  1. Convert to grayscale
  2. Upscale to 300dpi-equivalent if small
  3. Deskew (rotate correction)
  4. Adaptive threshold / denoise
  5. Tesseract with best config
"""

import math
import os
import re
import time

try:
    import fitz  # PyMuPDF

    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

try:
    import pytesseract
    from pytesseract import Output

    HAS_TESSERACT = True
    try:
        pytesseract.get_tesseract_version()
        TESSERACT_BINARY_OK = True
    except Exception:
        TESSERACT_BINARY_OK = False
except ImportError:
    HAS_TESSERACT = False
    TESSERACT_BINARY_OK = False

try:
    from PIL import Image, ImageEnhance, ImageFilter, ImageOps

    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import cv2
    import numpy as np

    HAS_CV2 = True
except ImportError:
    HAS_CV2 = False

try:
    from unidecode import unidecode
except ImportError:

    def unidecode(s):
        return s


try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd

    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False


# Tesseract config — PSM 3 = fully automatic page segmentation
TESS_CONFIG = r"--oem 3 --psm 3"
# For single-column documents
TESS_CONFIG_SINGLE = r"--oem 3 --psm 6"
# For sparse text (receipts, forms)
TESS_CONFIG_SPARSE = r"--oem 3 --psm 11"

MIN_CHARS_PER_PAGE = 40  # below this → page is likely scanned, fall back to OCR


def check_dependencies() -> dict:
    return {
        "PyMuPDF": HAS_FITZ,
        "Tesseract binary": TESSERACT_BINARY_OK,
        "Pillow": HAS_PIL,
        "OpenCV": HAS_CV2,
        "python-docx": HAS_DOCX,
    }


# Image preprocessing
def _pil_to_cv2(pil_img):
    import numpy as np

    return cv2.cvtColor(np.array(pil_img.convert("RGB")), cv2.COLOR_RGB2BGR)


def _cv2_to_pil(cv2_img):
    return Image.fromarray(cv2.cvtColor(cv2_img, cv2.COLOR_BGR2RGB))


def _deskew(img_cv2):
    """Detect and correct skew angle using Hough transform."""
    gray = cv2.cvtColor(img_cv2, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    lines = cv2.HoughLinesP(
        edges, 1, math.pi / 180, threshold=100, minLineLength=100, maxLineGap=10
    )
    if lines is None:
        return img_cv2
    angles = []
    for line in lines:
        x1, y1, x2, y2 = line[0]
        if x2 != x1:
            angles.append(math.degrees(math.atan2(y2 - y1, x2 - x1)))
    if not angles:
        return img_cv2
    # Median angle, ignore near-vertical lines
    angles = [a for a in angles if abs(a) < 45]
    if not angles:
        return img_cv2
    median_angle = sorted(angles)[len(angles) // 2]
    if abs(median_angle) < 0.5:
        return img_cv2
    h, w = img_cv2.shape[:2]
    M = cv2.getRotationMatrix2D((w // 2, h // 2), median_angle, 1.0)
    return cv2.warpAffine(
        img_cv2, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE
    )


def _upscale_if_small(img: Image.Image, min_dpi_width: int = 1200) -> Image.Image:
    """Upscale small images — Tesseract accuracy drops below ~150dpi."""
    w, h = img.size
    if w < min_dpi_width:
        scale = min_dpi_width / w
        new_w, new_h = int(w * scale), int(h * scale)
        img = img.resize((new_w, new_h), Image.LANCZOS)
    return img


def preprocess_image(img: Image.Image) -> Image.Image:
    """
    Full preprocessing pipeline for OCR:
    grayscale → upscale → deskew → denoise → threshold
    """
    # Upscale if too small
    img = _upscale_if_small(img)

    if HAS_CV2:
        cv = _pil_to_cv2(img)

        # Deskew
        cv = _deskew(cv)

        # Convert to grayscale
        gray = cv2.cvtColor(cv, cv2.COLOR_BGR2GRAY)

        # Denoise
        gray = cv2.fastNlMeansDenoising(gray, h=10)

        # Adaptive threshold → binary image (handles uneven lighting)
        binary = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10
        )

        # Slight sharpening kernel
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
        sharpened = cv2.filter2D(binary, -1, kernel)

        return Image.fromarray(sharpened)
    else:
        # PIL-only fallback
        img = img.convert("L")  # grayscale
        img = ImageOps.autocontrast(img, cutoff=2)  # contrast stretch
        img = img.filter(ImageFilter.SHARPEN)
        return img


# Core OCR functions
def ocr_image_file(image_path: str) -> str:
    """OCR a single image file with preprocessing."""
    if not HAS_TESSERACT or not TESSERACT_BINARY_OK:
        raise RuntimeError(
            "Tesseract not available. Install tesseract-ocr system package."
        )
    if not HAS_PIL:
        raise RuntimeError("Pillow not installed.")

    img = Image.open(image_path).convert("RGB")
    processed = preprocess_image(img)

    # Try multiple PSM configs, pick the one with most text
    results = []
    for config in [TESS_CONFIG, TESS_CONFIG_SINGLE, TESS_CONFIG_SPARSE]:
        try:
            text = pytesseract.image_to_string(processed, config=config, lang="eng")
            results.append(text)
        except Exception:
            pass

    if not results:
        raise RuntimeError("Tesseract failed on this image.")

    # Return longest result (most text extracted)
    best = max(results, key=lambda t: len(t.strip()))
    return best.strip()


def _pdf_page_to_image(page) -> Image.Image:
    """Render a PDF page to PIL Image at 300dpi."""
    mat = fitz.Matrix(300 / 72, 300 / 72)  # 300dpi
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img


def extract_pdf(pdf_path: str) -> tuple[str, int]:
    """
    Smart PDF extraction:
    - Try native text first (fast)
    - If page has <MIN_CHARS_PER_PAGE → render page and OCR it (scanned PDF)
    Returns (full_text, page_count)
    """
    if not HAS_FITZ:
        raise RuntimeError("PyMuPDF not installed. Run: pip install PyMuPDF")

    doc = fitz.open(pdf_path)
    page_texts = []
    ocr_pages = 0

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        native_text = page.get_text().strip()

        if len(native_text) >= MIN_CHARS_PER_PAGE:
            # Good native text
            page_texts.append(unidecode(native_text))
        elif HAS_TESSERACT and TESSERACT_BINARY_OK:
            # Scanned page — render and OCR
            img = _pdf_page_to_image(page)
            processed = preprocess_image(img)
            try:
                ocr_text = pytesseract.image_to_string(
                    processed, config=TESS_CONFIG, lang="eng"
                )
                page_texts.append(ocr_text.strip())
                ocr_pages += 1
            except Exception as e:
                page_texts.append(f"[Page {page_num + 1} OCR failed: {e}]")
        else:
            page_texts.append(
                native_text or f"[Page {page_num + 1}: no text, Tesseract unavailable]"
            )

    doc.close()
    full_text = "\n\n--- Page Break ---\n\n".join(t for t in page_texts if t)
    note = f"\n\n[{ocr_pages} page(s) processed via OCR]" if ocr_pages else ""
    return full_text + note, len(page_texts)


def extract_pdf_sections(pdf_path: str) -> list[dict]:
    """Extract heading/content pairs using font analysis."""
    if not HAS_FITZ:
        raise RuntimeError("PyMuPDF not installed.")

    doc = fitz.open(pdf_path)
    rows = []
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    text = unidecode(span["text"]).strip()
                    if text:
                        rows.append(
                            {
                                "text": text,
                                "size": span["size"],
                                "bold": "bold" in span["font"].lower(),
                            }
                        )
    doc.close()

    if not rows:
        return [{"heading": "Full Text", "content": "No text found"}]

    sizes = [r["size"] for r in rows]
    body_size = max(set(sizes), key=sizes.count)
    sections, current_heading, current_body = [], "Document", []

    for row in rows:
        is_heading = row["size"] > body_size + 1.5 or (
            row["bold"] and row["size"] >= body_size + 0.5
        )
        if is_heading:
            if current_body:
                sections.append(
                    {"heading": current_heading, "content": " ".join(current_body)}
                )
            current_heading = row["text"]
            current_body = []
        else:
            current_body.append(row["text"])

    if current_body:
        sections.append({"heading": current_heading, "content": " ".join(current_body)})

    return sections or [
        {"heading": "Full Text", "content": " ".join(r["text"] for r in rows)}
    ]


def sections_to_word(sections: list[dict], output_path: str) -> str:
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed.")
    doc = DocxDocument()
    for sec in sections:
        doc.add_heading(sec["heading"], level=1)
        p = doc.add_paragraph(sec["content"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        doc.add_paragraph("")
    doc.save(output_path)
    return output_path


def summarize_text(text: str, ratio: float = 0.3) -> str:
    """Extractive summary by sentence TF scoring."""
    sentences = [s.strip() for s in re.split(r"[.!?\n]+", text) if len(s.strip()) > 20]
    if not sentences:
        return text[:3000]
    words = re.findall(r"\w+", text.lower())
    freq: dict = {}
    for w in words:
        if len(w) > 3:
            freq[w] = freq.get(w, 0) + 1
    scored = sorted(
        [
            (sum(freq.get(w.lower(), 0) for w in re.findall(r"\w+", s)), s)
            for s in sentences
        ],
        key=lambda x: -x[0],
    )
    keep = max(3, int(len(scored) * ratio))
    return ". ".join(s for _, s in scored[:keep]) + "."


def answer_question(question: str, text: str) -> str:
    """Keyword-match sentence retrieval."""
    keywords = [w for w in re.findall(r"\w+", question.lower()) if len(w) > 3]
    if not keywords:
        return "Please ask a more specific question."
    sentences = [s.strip() for s in re.split(r"[.!?\n]+", text) if len(s.strip()) > 15]
    scored = sorted(
        [
            (sum(1 for kw in keywords if kw in s.lower()), s)
            for s in sentences
            if any(kw in s.lower() for kw in keywords)
        ],
        key=lambda x: -x[0],
    )
    return (
        " ".join(s for _, s in scored[:5])
        if scored
        else "No relevant information found."
    )


# Main dispatcher
def process_job(job_type: str, file_path: str, extra: dict = None) -> dict:
    """
    Run OCR job synchronously (called in thread executor from async context).
    Always returns a result dict — never raises.
    """
    start = time.time()
    extra = extra or {}
    result = {
        "text": None,
        "file_path": None,
        "error": None,
        "page_count": None,
        "processing_time_ms": 0,
    }

    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found on disk: {file_path}")

        if job_type == "ocr_image":
            result["text"] = ocr_image_file(file_path)

        elif job_type == "pdf_extract":
            text, pages = extract_pdf(file_path)
            result["text"] = text or "(No text extracted)"
            result["page_count"] = pages

        elif job_type == "pdf_summarize":
            text, pages = extract_pdf(file_path)
            result["text"] = summarize_text(text, float(extra.get("ratio", 0.3)))
            result["page_count"] = pages

        elif job_type == "pdf_to_word":
            sections = extract_pdf_sections(file_path)
            out_path = file_path.rsplit(".", 1)[0] + "_extracted.docx"
            sections_to_word(sections, out_path)
            result["file_path"] = out_path
            result["text"] = f"Extracted {len(sections)} section(s)"
            result["page_count"] = len(sections)

        elif job_type == "pdf_qa":
            text, pages = extract_pdf(file_path)
            result["text"] = answer_question(extra.get("question", ""), text)
            result["page_count"] = pages

        elif job_type == "pdf_to_markdown":
            md_text, pages = pdf_to_markdown(file_path)
            out_path = file_path.rsplit(".", 1)[0] + ".md"
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(md_text)
            result["text"] = md_text[:2000] + ("..." if len(md_text) > 2000 else "")
            result["file_path"] = out_path
            result["page_count"] = pages

        elif job_type == "pdf_merge":
            # extra["input_paths"] — list of additional local PDF paths to merge
            # The primary file_path is merged first
            input_paths = [file_path] + (extra.get("input_paths") or [])
            out_path = file_path.rsplit(".", 1)[0] + "_merged.pdf"
            _, total_pages = pdf_merge(input_paths, out_path)
            result["file_path"] = out_path
            result["text"] = (
                f"Merged {len(input_paths)} PDF(s) into {total_pages} pages"
            )
            result["page_count"] = total_pages

        elif job_type == "pdf_split":
            from_page = int(extra.get("from_page", 1))
            to_page = int(extra.get("to_page", 1))
            out_path = file_path.rsplit(".", 1)[0] + f"_pages_{from_page}-{to_page}.pdf"
            _, page_count = pdf_split(file_path, from_page, to_page, out_path)
            result["file_path"] = out_path
            result["text"] = (
                f"Extracted pages {from_page}–{to_page} ({page_count} page(s))"
            )
            result["page_count"] = page_count

        elif job_type == "pdf_compress":
            out_path = file_path.rsplit(".", 1)[0] + "_compressed.pdf"
            _, reduction_pct = pdf_compress(file_path, out_path)
            result["file_path"] = out_path
            result["text"] = f"Compressed PDF — {reduction_pct}% size reduction"
            result["page_count"] = 1

        elif job_type == "images_to_pdf":
            # extra["image_paths"] — additional image paths to include after file_path
            image_paths = [file_path] + (extra.get("image_paths") or [])
            out_path = file_path.rsplit(".", 1)[0] + "_combined.pdf"
            _, page_count = images_to_pdf(image_paths, out_path)
            result["file_path"] = out_path
            result["text"] = f"Combined {page_count} image(s) into PDF"
            result["page_count"] = page_count

        else:
            raise ValueError(f"Unknown job type: {job_type}")

    except Exception as e:
        result["error"] = f"{type(e).__name__}: {e}"

    result["processing_time_ms"] = int((time.time() - start) * 1000)
    return result


# Document Studio handlers
def pdf_to_markdown(pdf_path: str) -> tuple[str, int]:
    """
    Convert a PDF to Markdown format.
    Uses PyMuPDF's markdown output (available in PyMuPDF >= 1.24.0).
    Falls back to plain text with heading detection for older versions.
    """
    if not HAS_FITZ:
        raise RuntimeError("PyMuPDF not installed.")

    doc = fitz.open(pdf_path)
    pages_md = []

    for i, page in enumerate(doc):
        # Try native markdown output first (PyMuPDF 1.24+)
        try:
            md = page.get_text("markdown")
            if md and md.strip():
                pages_md.append(md.strip())
                continue
        except Exception:
            pass

        # Fallback: simulate markdown from block structure
        lines = []
        for block in page.get_text("dict")["blocks"]:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    text = unidecode(span["text"]).strip()
                    if not text:
                        continue
                    size = span["size"]
                    bold = "bold" in span["font"].lower()
                    if size >= 18:
                        lines.append(f"\n# {text}")
                    elif size >= 14 or (bold and size >= 12):
                        lines.append(f"\n## {text}")
                    elif bold:
                        lines.append(f"\n**{text}**")
                    else:
                        lines.append(text)
        pages_md.append("\n".join(lines))

    doc.close()
    full_md = "\n\n---\n\n".join(
        f"<!-- Page {i+1} -->\n{md}" for i, md in enumerate(pages_md) if md.strip()
    )
    return full_md, len(pages_md)


def pdf_merge(input_paths: list[str], output_path: str) -> tuple[str, int]:
    """
    Merge multiple PDFs into a single output file.
    input_paths — list of local PDF file paths in desired order.
    Returns (output_path, total_page_count).
    """
    if not HAS_FITZ:
        raise RuntimeError("PyMuPDF not installed.")
    if len(input_paths) < 2:
        raise ValueError("At least two PDFs required for merge.")

    merged = fitz.open()
    total_pages = 0

    for path in input_paths:
        src = fitz.open(path)
        merged.insert_pdf(src)
        total_pages += len(src)
        src.close()

    merged.save(output_path, deflate=True, garbage=3)
    merged.close()
    return output_path, total_pages


def pdf_split(
    pdf_path: str, from_page: int, to_page: int, output_path: str
) -> tuple[str, int]:
    """
    Extract a page range from a PDF (1-indexed, inclusive).
    Returns (output_path, page_count).
    """
    if not HAS_FITZ:
        raise RuntimeError("PyMuPDF not installed.")

    src = fitz.open(pdf_path)
    total = len(src)

    from_page = max(1, min(from_page, total))
    to_page = max(from_page, min(to_page, total))

    out = fitz.open()
    out.insert_pdf(src, from_page=from_page - 1, to_page=to_page - 1)  # 0-indexed
    out.save(output_path, deflate=True, garbage=3)
    out.close()
    src.close()

    page_count = to_page - from_page + 1
    return output_path, page_count


def pdf_compress(pdf_path: str, output_path: str) -> tuple[str, int]:
    """
    Reduce PDF file size using PyMuPDF's garbage collection + deflate compression.
    garbage=4 removes redundant objects; deflate=True uses zlib on streams.
    """
    if not HAS_FITZ:
        raise RuntimeError("PyMuPDF not installed.")

    doc = fitz.open(pdf_path)
    doc.save(
        output_path,
        garbage=4,  # aggressive cross-reference + object deduplication
        deflate=True,  # compress all streams
        clean=True,  # clean content streams
        deflate_images=True,
        deflate_fonts=True,
    )
    doc.close()

    original_size = os.path.getsize(pdf_path)
    compressed_size = os.path.getsize(output_path)
    reduction_pct = (
        round((1 - compressed_size / original_size) * 100, 1) if original_size else 0
    )

    return output_path, reduction_pct  # returns output_path + reduction %


def images_to_pdf(image_paths: list[str], output_path: str) -> tuple[str, int]:
    """
    Combine one or more images into a single PDF.
    Each image becomes one page.
    Returns (output_path, page_count).
    """
    if not HAS_FITZ:
        raise RuntimeError("PyMuPDF not installed.")
    if not image_paths:
        raise ValueError("No images provided.")

    doc = fitz.open()

    for img_path in image_paths:
        img_doc = fitz.open(img_path)  # works for JPEG, PNG, TIFF, BMP
        rect = img_doc[0].rect
        pdf_bytes = img_doc.convert_to_pdf()  # convert image to single-page PDF
        img_doc.close()

        img_pdf = fitz.open("pdf", pdf_bytes)
        doc.insert_pdf(img_pdf)
        img_pdf.close()

    doc.save(output_path, deflate=True)
    doc.close()
    return output_path, len(image_paths)
